"""Research-report engine (Phase C).

Every research skill shares ONE output contract: a polished ``.docx`` deep-dive
PLUS a machine-readable ``.json`` sidecar, both saved to

    4_Areas/Investing/Stock_Reports/<TARGET>/<YYYY-MM-DD>_<TARGET>[__skill].{docx,json}

The dashboard reads the sidecar (never parses Word). This module owns the whole
pipeline: build the prompt from a :class:`~modules.agents.skills.research.skills.ResearchSkill`,
run ``claude -p`` with web research pre-authorized, guard against refusals /
training-knowledge-only output (same contract as the idea-validator runner —
nothing is written on a bad run), extract + enrich the sidecar, then write both
files via Python file I/O (vault-write rule) with ``.bak`` backups.
"""
from __future__ import annotations

import json
import re
import shutil
from dataclasses import dataclass
from datetime import date, datetime
from pathlib import Path

from core.config import STOCK_REPORTS_DIR, get_logger
from modules.agents import data_sources
from modules.agents.claude_cli import run_claude

logger = get_logger(__name__)

# Pre-authorize the research tools — headless `claude -p` cannot answer a
# permission prompt, so an un-allowed tool is silently denied and the report
# degrades to caveated training-data output (idea-validator Item #78).
RESEARCH_ALLOWED_TOOLS = ["Read", "Write", "Bash", "WebSearch", "WebFetch", "Glob", "Grep"]

# Marker the model emits right before the machine-readable sidecar block. Kept
# out of the .docx body (stripped before Word generation).
SIDECAR_MARKER = "<!-- SIDECAR_JSON -->"


@dataclass(frozen=True)
class ResearchSkill:
    """One research skill: persona + report outline + sidecar schema-by-example."""

    key: str                 # SKILL_REGISTRY name, e.g. "stock_analysis"
    label: str               # display name
    target_kind: str         # "ticker" | "ticker_or_fund" | "sector" | "none"
    persona: str             # analyst framing line(s)
    outline: str             # the markdown section structure the report must follow
    sidecar_example: dict     # example sidecar object = the schema, shown to the model
    timeout: int = 900


# --- refusal / quality guards (mirror idea_validator runner) -----------------
_REFUSAL_MARKERS = (
    "i won't fabricate", "won't fabricate", "i can't fabricate", "cannot fabricate",
    "i can't help", "i cannot help", "i need more", "which ticker do you",
)
_NO_WEBSEARCH_MARKERS = (
    "no web search", "web search was unavailable", "web search unavailable",
    "websearch` permission denied", "websearch permission denied",
    "training data only", "training knowledge", "without access to current",
    "model cutoff",
)


def _looks_like_refusal(output: str) -> bool:
    text = output.strip()
    if not text:
        return True
    low = text.lower()
    if any(m in low for m in _REFUSAL_MARKERS):
        return True
    # A real report is long, structured Markdown. A short, heading-free reply is a chat refusal.
    if len(text) < 1500 and "\n## " not in text and not text.startswith("## "):
        return True
    return False


def _looks_like_no_websearch(output: str) -> bool:
    low = output.lower()
    return any(m in low for m in _NO_WEBSEARCH_MARKERS)


# --- target / path helpers ---------------------------------------------------
def safe_target(target: str | None) -> str:
    """Normalize a target into a filesystem- and ticker-safe token.

    Tickers/funds are upper-cased; free-text sectors collapse to an upper
    SNAKE token. Empty target (macro brief) -> ``MACRO``.
    """
    raw = (target or "").strip()
    if not raw:
        return "MACRO"
    token = re.sub(r"[^A-Za-z0-9._-]+", "_", raw).strip("_")
    return (token.upper() or "MACRO")[:40]


def report_basename(skill: ResearchSkill, target_token: str, on: date) -> str:
    """``<DATE>_<TARGET>`` for stock-analysis (exact v3 spec), suffixed elsewhere.

    The skill-key suffix keeps same-day reports for one ticker (e.g. a stock
    analysis AND a short thesis on BABA) from overwriting each other.
    """
    base = f"{on.isoformat()}_{target_token}"
    return base if skill.key == "stock_analysis" else f"{base}__{skill.key}"


def report_paths(skill: ResearchSkill, target_token: str, on: date) -> tuple[Path, Path]:
    """Absolute (docx, json) paths under the ticker's report folder."""
    folder = STOCK_REPORTS_DIR / target_token
    base = report_basename(skill, target_token, on)
    return folder / f"{base}.docx", folder / f"{base}.json"


def _vault_relative(p: Path) -> str:
    """Vault-relative POSIX path string for the sidecar ``report_path`` field."""
    from core.config import VAULT_PATH

    try:
        return p.relative_to(VAULT_PATH).as_posix()
    except ValueError:
        return p.as_posix()


# --- prompt assembly ---------------------------------------------------------
def _ticker_context(target_token: str) -> str:
    """Live price + headlines block for ticker-shaped targets (grounds the model)."""
    snap = data_sources.price_snapshot(target_token)
    if snap.get("error"):
        return f"(no live price for {target_token}: {snap['error']})"
    news = data_sources.combined_news(target_token, limit=8)
    lines = [
        f"Last close: {snap.get('last_close')} {snap.get('currency')}",
        f"1w change: {snap.get('week_change_pct'):.1f}%  ·  ~1m change: {snap.get('month_change_pct'):.1f}%",
        "",
        "Recent headlines:",
    ]
    for n in news:
        lines.append(f"- {n.get('title','')} ({n.get('publisher','')})")
    return "\n".join(lines)


def _build_prompt(skill: ResearchSkill, target: str | None, target_token: str, extra: str) -> str:
    today = date.today().isoformat()
    context = ""
    if skill.target_kind in {"ticker", "ticker_or_fund"}:
        context = _ticker_context(target_token)
    target_line = (
        "Target: (none — broad market read)"
        if skill.target_kind == "none"
        else f"Target: {target or target_token}"
    )
    sidecar_schema = json.dumps(skill.sidecar_example, indent=2, ensure_ascii=False)
    return f"""{skill.persona}

{target_line}
As-of date: {today}

# LIVE CONTEXT (ground every figure; verify with web search)
{context or '(no pre-fetched context — use web search)'}
{extra}

# REQUIRED REPORT STRUCTURE (Markdown)
{skill.outline}

Tone: dense senior-analyst prose. Every quantitative claim cited (use web
search heavily). No hedging filler. Mark any genuine assumption explicitly.

# MANDATORY MACHINE-READABLE SIDECAR
After the full Markdown report, output the line `{SIDECAR_MARKER}` on its own,
then a single fenced ```json block matching THIS schema exactly (same keys,
realistic values; use null where unknown, [] for empty lists):

```json
{sidecar_schema}
```

Output the Markdown report first, then the marker, then the JSON block. Nothing after the JSON.
"""


# --- output parsing ----------------------------------------------------------
def _split_report_and_sidecar(output: str) -> tuple[str, dict]:
    """Return (clean markdown body, parsed sidecar dict). Raises on no valid JSON."""
    # Prefer the JSON fence that follows our marker; fall back to the last json fence.
    body = output
    sidecar_text: str | None = None

    if SIDECAR_MARKER in output:
        body, _, tail = output.partition(SIDECAR_MARKER)
        m = re.search(r"```json\s*(.*?)```", tail, re.DOTALL)
        if m:
            sidecar_text = m.group(1)
    if sidecar_text is None:
        fences = re.findall(r"```json\s*(.*?)```", output, re.DOTALL)
        if fences:
            sidecar_text = fences[-1]
            # Drop the trailing fenced block from the docx body.
            body = output[: output.rfind("```json")]

    if sidecar_text is None:
        raise RuntimeError(
            "No JSON sidecar block found in the report output — refusing to write a "
            "report without its machine-readable sidecar. Re-run."
        )
    try:
        sidecar = json.loads(sidecar_text.strip())
    except json.JSONDecodeError as exc:
        raise RuntimeError(f"Sidecar JSON did not parse ({exc}). Nothing written.") from exc
    if not isinstance(sidecar, dict):
        raise RuntimeError("Sidecar JSON was not an object. Nothing written.")
    return body.strip(), sidecar


def _enrich_sidecar(sidecar: dict, skill: ResearchSkill, target: str | None,
                    target_token: str, docx_path: Path, on: date) -> dict:
    """Stamp the authoritative, server-known fields over whatever the model wrote.

    Guarantees the dashboard-critical keys (skill, target, date, report_path)
    are always present and correct regardless of model output.
    """
    sidecar["skill"] = skill.key
    sidecar["skill_label"] = skill.label
    sidecar.setdefault("ticker", target_token if skill.target_kind != "none" else None)
    sidecar["target"] = target or target_token
    sidecar["analysis_date"] = on.isoformat()
    sidecar["report_path"] = _vault_relative(docx_path)
    sidecar["generated_at"] = datetime.now().isoformat(timespec="seconds")
    sidecar.setdefault("sources_as_of", on.isoformat())
    # Backfill a live price for ticker reports if the model left it null.
    if skill.target_kind in {"ticker", "ticker_or_fund"} and not sidecar.get("current_price_at_analysis"):
        snap = data_sources.price_snapshot(target_token)
        if not snap.get("error"):
            sidecar["current_price_at_analysis"] = round(float(snap["last_close"]), 2)
            sidecar.setdefault("currency", snap.get("currency"))
    return sidecar


# --- docx generation ---------------------------------------------------------
def _md_to_docx(markdown: str, path: Path, title: str) -> None:
    """Render the markdown report to .docx (python-docx), with a title block."""
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=0)
    for line in markdown.splitlines():
        s = line.rstrip()
        if s.startswith("#### "):
            doc.add_heading(s[5:], level=3)
        elif s.startswith("### "):
            doc.add_heading(s[4:], level=2)
        elif s.startswith("## "):
            doc.add_heading(s[3:], level=1)
        elif s.startswith("# "):
            doc.add_heading(s[2:], level=1)
        elif s.lstrip().startswith(("- ", "* ")):
            doc.add_paragraph(s.lstrip()[2:], style="List Bullet")
        elif s.strip():
            doc.add_paragraph(s)
    path.parent.mkdir(parents=True, exist_ok=True)
    if path.exists():
        shutil.copy2(path, path.with_suffix(".docx.bak"))
    doc.save(path)


def _write_sidecar(sidecar: dict, path: Path) -> None:
    """Write the JSON sidecar via Python file I/O with a .bak backup (vault rule)."""
    path.parent.mkdir(parents=True, exist_ok=True)
    if path.exists():
        shutil.copy2(path, path.with_suffix(".json.bak"))
    path.write_text(json.dumps(sidecar, indent=2, ensure_ascii=False, default=str), encoding="utf-8")


# --- public entry point ------------------------------------------------------
def generate(skill: ResearchSkill, target: str | None = None, extra_context: str = "") -> dict:
    """Run one research skill end-to-end. Returns a small result dict.

    Raises (writing nothing) on a refusal, training-knowledge-only output, or a
    missing/invalid sidecar — same fail-loud contract as the validator runner.
    """
    if skill.target_kind != "none" and not (target or "").strip():
        raise ValueError(f"Skill '{skill.key}' requires a target ({skill.target_kind}).")

    target_token = safe_target(target)
    on = date.today()
    docx_path, json_path = report_paths(skill, target_token, on)

    prompt = _build_prompt(skill, target, target_token, extra_context)
    logger.info("research.generate skill=%s target=%s", skill.key, target_token)
    output = run_claude(prompt, timeout=skill.timeout, allowed_tools=RESEARCH_ALLOWED_TOOLS)

    if _looks_like_refusal(output):
        raise RuntimeError(
            f"{skill.label} for '{target_token}' returned a refusal/clarification "
            f"({len(output.strip())} chars), not a report. Nothing written."
        )
    if _looks_like_no_websearch(output):
        raise RuntimeError(
            f"{skill.label} for '{target_token}' looks like training-knowledge-only "
            "output (web search unavailable). Nothing written — check WebSearch, re-run."
        )

    body, sidecar = _split_report_and_sidecar(output)
    sidecar = _enrich_sidecar(sidecar, skill, target, target_token, docx_path, on)

    title = sidecar.get("title") or f"{skill.label} — {target or target_token} ({on.isoformat()})"
    _md_to_docx(body, docx_path, title)
    _write_sidecar(sidecar, json_path)

    # Stamp the declarative registry's last_run (best-effort; never fail the run).
    try:
        from modules.agents.skills.research import catalog

        catalog.stamp_last_run(skill.key, "completed", on=datetime.now())
    except Exception:  # noqa: BLE001
        logger.warning("could not stamp last_run for %s", skill.key, exc_info=True)

    logger.info("research report written: %s", docx_path)
    return {
        "skill": skill.key,
        "target": target_token,
        "docx": str(docx_path),
        "json": str(json_path),
        "report_path": sidecar["report_path"],
        "recommendation": sidecar.get("recommendation"),
    }


# --- reports library index (read side; the Phase D UI consumes this) ---------
def list_reports() -> list[dict]:
    """Every report sidecar on disk, newest first — powers the Reports library."""
    if not STOCK_REPORTS_DIR.exists():
        return []
    out: list[dict] = []
    for jf in STOCK_REPORTS_DIR.glob("*/*.json"):
        if jf.name.endswith((".json.bak", ".bak")):
            continue
        try:
            data = json.loads(jf.read_text(encoding="utf-8"))
        except (json.JSONDecodeError, OSError):
            continue
        data["sidecar_path"] = str(jf)
        out.append(data)
    return sorted(out, key=lambda d: (d.get("analysis_date", ""), d.get("generated_at", "")), reverse=True)


def latest_for_ticker(ticker: str) -> dict | None:
    """Latest stock-analysis sidecar for a ticker — the PositionPlanCard source."""
    token = safe_target(ticker)
    reports = [
        r for r in list_reports()
        if r.get("skill") == "stock_analysis" and safe_target(r.get("ticker") or r.get("target")) == token
    ]
    return reports[0] if reports else None
